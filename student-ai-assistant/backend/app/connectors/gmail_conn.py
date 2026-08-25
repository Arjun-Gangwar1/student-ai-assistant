"""
Gmail connector — structured email sync with attachment text extraction.

⚠ RESTRICTED SCOPE. `gmail.readonly` is classified restricted by Google.
Serving more than 100 lifetime users requires OAuth verification plus an annual
CASA security assessment (~$540/yr). Gated behind GMAIL_ENABLED and a per-student
consent flag so that obligation is a deliberate choice. See docs/PRIVACY.md.

Each message is written to two places:
  emails  — the structured record a student can browse and filter
  items   — the RAG layer: cleaned text, embedding, category

Attachments with extractable text get their own `items` row, so a question like
"what was in the PDF the prof sent about Assignment 3" can retrieve the PDF's
contents rather than just the covering email.
"""

import base64
import io
import logging
import re
from datetime import datetime, timedelta, timezone

from bs4 import BeautifulSoup
from starlette.concurrency import run_in_threadpool

from app.config import settings
from app.connectors.google_auth import GoogleAuthError, build_service, has_scope
from app.db import queries

logger = logging.getLogger(__name__)

DEFAULT_LOOKBACK_DAYS = 30
MAX_EMAILS_PER_SYNC = 50
MAX_ATTACHMENT_BYTES = 5 * 1024 * 1024
MAX_BODY_CHARS = 8000
MAX_EXTRACTED_CHARS = 6000

TEXT_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "text/plain",
    "text/markdown",
    "text/csv",
}
TEXT_EXTENSIONS = (".pdf", ".docx", ".doc", ".txt", ".md", ".csv")

DOC_LINK_RE = re.compile(
    r"https?://(?:docs\.google\.com|drive\.google\.com|github\.com|"
    r"classroom\.google\.com|forms\.gle|[a-z0-9-]+\.iitdh\.ac\.in)[^\s\"'>)\]]*",
    re.IGNORECASE,
)
SENDER_RE = re.compile(r'^(?:"?([^"<]*)"?\s*)?<([^>]+)>$')

# Unicode spaces and invisible characters. Mail from web composers is full of
# these; left in, they fragment tokens and wreck both search and embeddings.
UNICODE_SPACES = re.compile(r"[   -   　]")
INVISIBLE_CHARS = re.compile(r"[­​-‏‪-‮⁠﻿͏]")


# ─── Parsing helpers ─────────────────────────────────────────────────────────

def _parse_sender(raw: str) -> tuple[str, str]:
    """'Arjun <a@b.com>' → ('Arjun', 'a@b.com')"""
    raw = (raw or "").strip()
    if match := SENDER_RE.match(raw):
        name = (match.group(1) or "").strip().strip('"')
        email = (match.group(2) or "").strip().lower()
        return name or email.split("@")[0], email
    if "@" in raw:
        return raw.split("@")[0].replace(".", " ").title(), raw.lower()
    return raw, ""


def _get_header(headers: list[dict], name: str) -> str:
    lowered = name.lower()
    return next((h["value"] for h in headers if h.get("name", "").lower() == lowered), "")


def _decode_part(data: str) -> str:
    """Base64url-decode a Gmail body part, tolerating missing padding."""
    if not data:
        return ""
    try:
        cleaned = data.replace("-", "+").replace("_", "/").strip()
        padded = cleaned + "=" * (-len(cleaned) % 4)
        return base64.b64decode(padded).decode("utf-8", errors="replace")
    except Exception:
        return ""


def _html_to_text(html: str) -> str:
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "head", "meta", "link", "img"]):
        tag.decompose()
    # separator=" " prevents adjacent block elements merging into "wordword".
    return soup.get_text(separator=" ", strip=True)


def _extract_body_parts(payload: dict) -> tuple[str, str]:
    """Recursively collect (plain, html) from a MIME tree."""
    mime = payload.get("mimeType", "")

    if mime.startswith("text/plain"):
        return _decode_part(payload.get("body", {}).get("data", "")), ""
    if mime.startswith("text/html"):
        return "", _decode_part(payload.get("body", {}).get("data", ""))

    if mime.startswith("multipart"):
        plains, htmls = [], []
        for part in payload.get("parts", []):
            plain, html = _extract_body_parts(part)
            if plain:
                plains.append(plain)
            if html:
                htmls.append(html)
        return "\n".join(plains), "\n".join(htmls)

    return "", ""


def _strip_quoted_replies(text: str) -> str:
    """Drop '>' quoted blocks and 'On ... wrote:' trailers."""
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith(">"):
            break
        if re.match(r"^On .{10,120}\s+wrote:$", stripped):
            break
        if re.match(r"^-{2,}\s*Forwarded message\s*-{2,}$", stripped, re.I):
            break
        lines.append(line)
    return "\n".join(lines).strip()


def _normalise(text: str) -> str:
    text = UNICODE_SPACES.sub(" ", text)
    text = INVISIBLE_CHARS.sub("", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def clean_body(payload: dict, snippet: str) -> str:
    """Best available plain text for a message."""
    plain, html = _extract_body_parts(payload)
    stripped = plain.strip()

    # Some generators emit a Python bytes repr into text/plain; prefer the HTML
    # part when that happens.
    if stripped and not stripped.startswith(("b'", 'b"')):
        text = _strip_quoted_replies(plain)
    elif html.strip():
        text = _strip_quoted_replies(_html_to_text(html))
    else:
        text = snippet

    return _normalise(text)[:MAX_BODY_CHARS]


def extract_doc_links(body: str) -> list[str]:
    return list(dict.fromkeys(DOC_LINK_RE.findall(body)))


def _collect_attachments(payload: dict, out: list[dict]) -> None:
    if (filename := payload.get("filename")) and (
        attachment_id := payload.get("body", {}).get("attachmentId")
    ):
        out.append({
            "filename": filename,
            "mime_type": payload.get("mimeType", ""),
            "size_bytes": payload.get("body", {}).get("size", 0),
            "attachment_id": attachment_id,
        })
    for part in payload.get("parts", []):
        _collect_attachments(part, out)


def extract_text_from_bytes(
    data: bytes, mime_type: str, filename: str, max_chars: int = MAX_EXTRACTED_CHARS
) -> str:
    """Pull text out of a file. Returns '' when unsupported."""
    lower = filename.lower()
    try:
        if mime_type == "application/pdf" or lower.endswith(".pdf"):
            import pdfplumber

            with pdfplumber.open(io.BytesIO(data)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages[:20]]
            return _normalise("\n".join(pages))[:max_chars]

        if lower.endswith((".docx", ".doc")) or mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            from docx import Document

            document = Document(io.BytesIO(data))
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            # Tables carry timetables and exam schedules — the highest-value
            # content in a campus attachment.
            for table in document.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            return _normalise("\n".join(paragraphs))[:max_chars]

        if mime_type in ("text/plain", "text/markdown", "text/csv") or lower.endswith(
            (".txt", ".md", ".csv")
        ):
            return _normalise(data.decode("utf-8", errors="ignore"))[:max_chars]

    except Exception as exc:
        logger.warning("Text extraction failed (%s): %s", filename, exc)

    return ""


# ─── Sync ────────────────────────────────────────────────────────────────────

def gmail_allowed_for(student: dict) -> tuple[bool, str]:
    """
    Whether Gmail may be synced for this student, and why not if not.

    Three gates, all of which must pass: the deployment flag, an optional
    allowlist for piloting before verification, and the student's own consent.
    """
    if not settings.gmail_enabled:
        return False, "GMAIL_ENABLED is false"
    if not student.get("gmail_enabled"):
        return False, "student has not enabled Gmail"
    if not has_scope(student, "gmail"):
        return False, "gmail scope not granted"

    allowlist = settings.gmail_allowlist_emails
    if allowlist and (student.get("email") or "").lower() not in allowlist:
        return False, "not in GMAIL_ALLOWLIST"

    return True, ""


async def sync_student_gmail(student: dict, max_emails: int = MAX_EMAILS_PER_SYNC) -> dict:
    student_id = str(student["id"])
    counts = {"emails": 0, "attachments": 0, "errors": 0, "skipped": None}

    allowed, reason = gmail_allowed_for(student)
    if not allowed:
        counts["skipped"] = reason
        logger.debug("Gmail sync skipped for %s: %s", student_id, reason)
        return counts

    try:
        service = await build_service(student, "gmail", "v1")
    except GoogleAuthError as exc:
        logger.warning("Gmail auth failed for %s: %s", student_id, exc)
        counts["errors"] += 1
        return counts

    # Incremental: fetch only what arrived since the newest message held. The
    # previous version re-fetched a fixed 30-day window every three minutes,
    # re-downloading and re-parsing the same messages indefinitely.
    latest = await queries.latest_email_timestamp(student_id)
    if latest:
        after = latest - timedelta(hours=1)     # small overlap for clock skew
    else:
        after = datetime.now(timezone.utc) - timedelta(days=DEFAULT_LOOKBACK_DAYS)

    query = (
        f"after:{int(after.timestamp())} "
        "-category:promotions -category:social -in:spam -in:trash"
    )

    try:
        listing = await run_in_threadpool(
            lambda: service.users()
            .messages()
            .list(userId="me", q=query, maxResults=max_emails)
            .execute()
        )
        message_refs = listing.get("messages", [])
    except Exception as exc:
        logger.error("Gmail list failed for %s: %s", student_id, exc)
        counts["errors"] += 1
        return counts

    for ref in message_refs:
        try:
            message = await run_in_threadpool(
                lambda mid=ref["id"]: service.users()
                .messages()
                .get(userId="me", id=mid, format="full")
                .execute()
            )
            attachments = await _process_message(service, student_id, message)
            counts["emails"] += 1
            counts["attachments"] += attachments
        except Exception as exc:
            logger.warning("Processing message %s failed: %s", ref["id"], exc)
            counts["errors"] += 1

    logger.info("Gmail sync for %s: %s", student_id, counts)
    return counts


async def _process_message(service, student_id: str, message: dict) -> int:
    payload = message.get("payload", {})
    headers = payload.get("headers", [])
    message_id = message["id"]

    subject = _get_header(headers, "Subject") or "(no subject)"
    sender_name, sender_email = _parse_sender(_get_header(headers, "From"))
    raw_to = _get_header(headers, "To")
    recipients = [r.strip() for r in raw_to.split(",") if r.strip()] if raw_to else []

    label_ids = message.get("labelIds", [])
    snippet = message.get("snippet", "")
    received_at = datetime.fromtimestamp(
        int(message.get("internalDate", 0)) / 1000, tz=timezone.utc
    )

    body_text = clean_body(payload, snippet)
    doc_links = extract_doc_links(body_text)

    attachment_parts: list[dict] = []
    _collect_attachments(payload, attachment_parts)

    # 1. items row — what RAG retrieves against
    item = await queries.upsert_item(
        student_id=student_id,
        source="gmail",
        source_id=message_id,
        raw_content=f"From: {sender_name} <{sender_email}>\nSubject: {subject}\n\n{body_text}",
        title=subject,
        metadata={
            "sender_name": sender_name,
            "sender_email": sender_email,
            "thread_id": message.get("threadId", ""),
            "labels": label_ids,
            "has_attachments": bool(attachment_parts),
            "doc_links": doc_links[:5],
        },
    )
    item_id = str(item["id"])

    # 2. emails row — what the inbox UI browses
    email_row = await queries.upsert_email(
        student_id=student_id,
        item_id=item_id,
        message_id=message_id,
        thread_id=message.get("threadId", ""),
        subject=subject,
        sender_name=sender_name,
        sender_email=sender_email,
        recipients=recipients,
        received_at=received_at,
        body_text=body_text,
        snippet=snippet,
        labels=label_ids,
        has_attachments=bool(attachment_parts),
        attachment_count=len(attachment_parts) + len(doc_links),
        is_read="UNREAD" not in label_ids,
    )
    email_id = str(email_row["id"])

    processed = 0
    for part in attachment_parts:
        try:
            await _process_attachment(
                service, student_id, email_id, message_id, subject, sender_name, part
            )
            processed += 1
        except Exception as exc:
            logger.warning("Attachment %s failed: %s", part.get("filename"), exc)

    # 3. Document links recorded as link-type attachments
    for url in doc_links:
        await queries.upsert_email_attachment(
            email_id=email_id,
            student_id=student_id,
            filename=(url.rstrip("/").split("/")[-1][:80] or "link"),
            mime_type="text/uri-list",
            attachment_type="link",
            url=url,
            extracted_text=f"Link shared in email {subject!r}: {url}",
        )

    return processed


async def _process_attachment(
    service, student_id, email_id, message_id, subject, sender_name, part
) -> None:
    filename = part["filename"]
    mime_type = part["mime_type"]
    size_bytes = part["size_bytes"]

    is_image = mime_type.startswith("image/")
    can_extract = mime_type in TEXT_MIME_TYPES or filename.lower().endswith(TEXT_EXTENSIONS)

    extracted_text = ""
    if can_extract and 0 < size_bytes <= MAX_ATTACHMENT_BYTES:
        blob = await run_in_threadpool(
            lambda: service.users()
            .messages()
            .attachments()
            .get(userId="me", messageId=message_id, id=part["attachment_id"])
            .execute()
        )
        raw = blob.get("data", "")
        padded = raw.replace("-", "+").replace("_", "/")
        padded += "=" * (-len(padded) % 4)
        extracted_text = extract_text_from_bytes(
            base64.b64decode(padded), mime_type, filename
        )

    attachment = await queries.upsert_email_attachment(
        email_id=email_id,
        student_id=student_id,
        filename=filename,
        mime_type=mime_type,
        size_bytes=size_bytes,
        extracted_text=extracted_text,
        attachment_type="image" if is_image else "file",
    )

    # Only attachments with real text become searchable items — an items row
    # with no content is retrieval noise.
    if extracted_text.strip():
        item = await queries.upsert_item(
            student_id=student_id,
            source="gmail_attachment",
            source_id=f"{message_id}:{filename}",
            raw_content=extracted_text,
            title=f"[Attachment] {filename}",
            metadata={
                "email_id": email_id,
                "filename": filename,
                "mime_type": mime_type,
                "sender_name": sender_name,
                "parent_subject": subject,
            },
        )
        await queries.link_attachment_item(str(attachment["id"]), str(item["id"]))
